# -*- coding: utf-8 -*-
"""Build the final docx report. All Chinese quotes use “ ”."""
import json, os, sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Load content
with open(r'E:\wuyi\数学建模半自动\research-assistant\outputs\content.json', 'r', encoding='utf-8') as f:
    C = json.load(f)

LQ = '“'  # "
RQ = '”'  # "

OUT = r'C:\Users\lenovo\Desktop\school\专业实践\思政'
IMG = r'E:\wuyi\数学建模半自动\research-assistant\outputs\思政论文配图'

doc = Document()

# Page setup
for sec in doc.sections:
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = Pt(20)

def set_font(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold

def add_p(text='', align=None, indent=False, spacing=20):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(spacing)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    return p

def add_run(p, text, name='宋体', size=12, bold=False):
    run = p.add_run(text)
    set_font(run, name, size, bold)
    return run

def add_body(text):
    p = add_p(indent=True)
    add_run(p, text)
    return p

def add_heading1(text):
    p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER, spacing=22)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, '黑体', 14, True)
    return p

def add_heading2(text):
    p = add_p(spacing=22)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, '黑体', 13, True)
    return p

def add_heading3(text):
    p = add_p(spacing=22)
    add_run(p, text, '黑体', 12, True)
    return p

def add_image(img_name, caption, width=14.0):
    path = os.path.join(IMG, img_name)
    if os.path.exists(path):
        p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Pt(8)
        p.add_run().add_picture(path, width=Cm(width))
        p2 = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.line_spacing = Pt(16)
        add_run(p2, caption, '楷体_GB2312', 10.5)

def add_table(headers, rows, caption, col_widths=None):
    p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    add_run(p, caption, '楷体_GB2312', 10.5)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        set_font(r, '宋体', 10.5, True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(ct))
            set_font(r, '宋体', 10.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def add_ref(text):
    p = add_p(indent=True)
    add_run(p, text)
    return p

# ===== COVER =====
for _ in range(6):
    add_p()
p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p, '思想政治理论课暑期社会实践', '黑体', 22, True)
p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p, '调 查 报 告', '黑体', 26, True)
for _ in range(4):
    add_p()

p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
p.paragraph_format.line_spacing = Pt(28)
add_run(p, '题目：' + C['title_main'], '宋体', 16, True)
p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
p.paragraph_format.line_spacing = Pt(28)
add_run(p, C['title_sub'], '宋体', 16, True)

for _ in range(6):
    add_p()

for line in ['学生姓名            _______________', '学    号            _______________',
             '指导教师            _______________', '学    院            _______________',
             '专业名称            _______________', '班    级            _______________']:
    p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.line_spacing = Pt(28)
    add_run(p, line, '宋体', 16)

for _ in range(2):
    add_p()
p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
p.paragraph_format.line_spacing = Pt(28)
add_run(p, '2026年8月', '宋体', 16)

doc.add_page_break()

# ===== AUTHORIZATION =====
p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(20)
add_run(p, '附件二\n实地调研资料授权书', '黑体', 16, True)

text = (
    '应中华人民共和国教育部思政司关于思政课社会实践要求，浙江财经大学马克思主义学院'
    '将于2026年暑期开展' + LQ + '沿着总书记的足迹，见证新时代伟大变革' + RQ + '实践活动。\n\n'
    '1. 同意项目组对受访人的访谈音像资料进行整理；整理的文稿经受访人校阅无误后签字确认。\n'
    '2. 受访人同意项目组发表或出版经受访人确认的访谈文稿，并使用受访人肖像。\n'
    '3. 受访人的访谈资料和肖像只用于学术研究和公益宣传。\n\n\n'
)
p = add_p(indent=True)
add_run(p, text)

for _ in range(3):
    add_p()
p = add_p(align=WD_ALIGN_PARAGRAPH.RIGHT)
add_run(p, '受访人（签名）：\n日期：')

doc.add_page_break()

# ===== TITLE PAGE =====
p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = Pt(30)
add_run(p, C['title_main'], '黑体', 16, True)
p = add_p(align=WD_ALIGN_PARAGRAPH.CENTER)
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.line_spacing = Pt(24)
add_run(p, '——' + C['title_sub'], '楷体_GB2312', 14)

# ===== ABSTRACT =====
p = add_p(indent=True)
add_run(p, '摘要：', '黑体', 12)
add_run(p, C['abstract'], '楷体_GB2312', 12)

p = add_p(indent=True)
add_run(p, '关键词：', '黑体', 12)
add_run(p, C['keywords'], '楷体_GB2312', 12)

add_p()

# ===== BODY SECTIONS =====
for sec in C['sections']:
    add_heading2(sec['heading'])

    if 'paras' in sec:
        for para_text in sec['paras']:
            add_body(para_text)

    if 'subsections' in sec:
        for sub in sec['subsections']:
            add_heading3(sub['heading'])
            for para_text in sub['paras']:
                add_body(para_text)
            # Insert images within subsections
            if '（二）状态' in sub['heading']:
                add_image('fig3_rd_gdp.png', '图3  丽水市莲都区GDP与R&D经费投入变化趋势（2006—2024年）', 13.0)
            elif '（三）响应' in sub['heading']:
                add_image('fig4_theil_index.png', '图4  丽水市莲都区城乡收入泰尔指数变化趋势（2006—2024年）', 13.0)

    # Insert images at section level
    if '一、引言' in sec['heading']:
        add_image('fig1_policy_timeline.png', '图1  山海协作政策演进关键节点（2002—2025年）')
    elif '二、调研设计' in sec['heading']:
        add_table(C['table1_headers'], C['table1_rows'], C['table1_caption'], [2.5, 1.5, 5.0, 5.5])
        add_image('fig2_geo_map.png', '图2  浙江省丽水市莲都区调研区位与点位空间分布图（数据来源：自然资源部标准地图服务，审图号：GS(2024)0650号；调研点位为团队实地标注）', 15.5)

# ===== REFERENCES =====
add_heading2('参考文献')
for ref in C['references']:
    add_ref(ref)

# ===== APPENDIX =====
doc.add_page_break()
add_heading2('附录一  访谈对象基本信息表')
add_table(C['appendix_headers'], C['appendix_rows'],
          '附表1  实地访谈对象基本信息', [2.8, 2.0, 4.5, 4.0, 2.0])

p = add_p(indent=True)
add_run(p, '注：应部分受访者要求，姓名均以姓氏+某（化名）形式呈现。访谈原始录音与文字记录另行归档。', '宋体', 10.5)

add_p()
add_heading2('附录二  评阅单')
add_p('（此处留空，由指导教师填写调查报告评阅单。）', indent=True)

# ===== SAVE =====
out_path = os.path.join(OUT, '跨越山海的共富图景——山海协作20年调研报告_v3_含地理图.docx')
doc.save(out_path)
print('Report saved to: ' + out_path)
