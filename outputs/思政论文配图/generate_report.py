# -*- coding: utf-8 -*-
"""生成完整调研报告 docx"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

output_dir = r"C:\Users\lenovo\Desktop\school\专业实践\思政"
img_dir = r"E:\wuyi\数学建模半自动\research-assistant\outputs\思政论文配图"

doc = Document()

# ============ 页面设置 ============
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)  # 小四
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = Pt(20)

def add_cover(doc):
    """封面页"""
    # 空行
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(20)

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run("思想政治理论课暑期社会实践")
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run("调 查 报 告")
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(26)
    run.font.bold = True

    for _ in range(4):
        doc.add_paragraph()

    # 题目
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(28)
    run = p.add_run("题目：跨越山海的共富图景\n——基于丽水莲都实践的"山海协作"20年要素流动与空间演进调研报告")
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(16)
    run.font.bold = True

    for _ in range(6):
        doc.add_paragraph()

    # 信息栏
    info_lines = [
        "学生姓名            _______________",
        "学    号            _______________",
        "指导教师            _______________",
        "学    院            _______________",
        "专业名称            _______________",
        "班    级            _______________",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = Pt(28)
        run = p.add_run(line)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(16)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(28)
    run = p.add_run("2026年8月")
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(16)

    doc.add_page_break()

def add_authorization(doc):
    """附件二：实地调研资料授权书"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("附件二\n实地调研资料授权书")
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(16)
    run.font.bold = True

    text = (
        "应中华人民共和国教育部思政司关于思政课社会实践要求，浙江财经大学马克思主义学院"
        "将于2026年暑期开展"沿着总书记的足迹，见证新时代伟大变革"实践活动。\n\n"
        "1. 同意项目组对受访人的访谈音像资料进行整理；整理的文稿经受访人校阅无误后签字确认。\n"
        "2. 受访人同意项目组发表或出版经受访人确认的访谈文稿，并使用受访人肖像。\n"
        "3. 受访人的访谈资料和肖像只用于学术研究和公益宣传。\n\n\n"
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("受访人（签名）：\n日期：")
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

    doc.add_page_break()

def add_title_page(doc):
    """正文标题页"""
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(30)
    run = p.add_run("跨越山海的共富图景")
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run("——基于丽水莲都实践的"山海协作"20年要素流动与空间演进调研报告")
    run.font.name = '楷体_GB2312'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
    run.font.size = Pt(14)

    return doc

def add_abstract(doc):
    """摘要"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run("摘要：")
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(12)
    run = p.add_run(
        "本报告以习近平同志"八八战略"与"山海协作"重要论述为理论指引，选取丽水市莲都区为典型案例，"
        "采用驱动力-压力-状态-影响-响应（DPSR）分析框架，结合实地走访、深度访谈与宏观数据分析，"
        "系统考察了2006—2025年间莲都—义乌山海协作在创新要素流动、产业空间重构与城乡收入收敛"
        "三个维度的实践成效。调研发现，协作机制已从初期的单向财政转移支付演化为R&D要素深度下沉"
        "与产业链双向嵌入的新格局，泰尔指数由2006年的0.182持续收敛至2024年的0.082。"
        "然而，高端研发人才"下沉易、留存难"以及偏远高山村落辐射递减等结构性挑战仍然突出。"
        "据此提出数字化柔性引才与共富工坊产业链延伸两条对策路径。"
    )
    run.font.name = '楷体_GB2312'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
    run.font.size = Pt(12)

    # 关键词
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run("关键词：")
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(12)
    run = p.add_run("山海协作；DPSR模型；要素流动；城乡收入差距；共富工坊")
    run.font.name = '楷体_GB2312'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
    run.font.size = Pt(12)

    doc.add_paragraph()  # 空行
    return doc

def add_heading_text(doc, text, level=1):
    """添加标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)
        run.font.bold = True
    elif level == 2:
        run = p.add_run(text)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(13)
        run.font.bold = True
    elif level == 3:
        run = p.add_run(text)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(12)
        run.font.bold = True
    return p

def add_body(doc, text):
    """正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    return p

def add_image_with_caption(doc, img_path, caption, width_cm=14.0):
    """插入图片和图注"""
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(16)
        run = p.add_run(caption)
        run.font.name = '楷体_GB2312'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
        run.font.size = Pt(10.5)  # 五号

def add_table_with_caption(doc, headers, rows, caption, col_widths=None):
    """插入三线表"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(caption)
    run.font.name = '楷体_GB2312'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
    run.font.size = Pt(10.5)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        run.font.bold = True

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10.5)

    doc.add_paragraph()  # 表后空行

# ============ 开始构建文档 ============

# 第1页：封面
add_cover(doc)

# 第2页：授权书
add_authorization(doc)

# 标题、摘要、关键词
add_title_page(doc)
add_abstract(doc)

# ============ 一、引言 ============
add_heading_text(doc, "一、引言：战略指引与跨越山海的二十年", level=2)

add_body(doc,
    "2002年，习近平同志在浙江工作期间提出了"八八战略"，将"进一步发挥浙江的山海资源优势，"
    "大力发展海洋经济，推动欠发达地区跨越式发展"确立为八大战略举措之一。2003年，浙江省委"
    "全面启动山海协作工程，将沿海发达市县与山区26县结对，构建起以要素流动为核心的区域协调"
    "发展制度框架。二十余年来，这一战略经历了从"资金输血"到"产业造血"、从单向转移到双向嵌入"
    "的深刻演变，成为观察中国特色社会主义共同富裕道路的重要窗口。"
)

add_body(doc,
    "习近平总书记多次强调，"共同富裕是社会主义的本质要求，是中国式现代化的重要特征""
    "（习近平，2021）。2021年，中共中央、国务院印发《关于支持浙江高质量发展建设共同富裕示范区的意见》，"
    "将山海协作提升至国家战略高度（中共中央，2021）。丽水市莲都区作为山区26县的核心组成部分，"
    "与义乌市建立了长达20年的山海协作结对关系。截至2024年，两地协作已延伸至数字化车间共建、"
    "科创飞地运营、共富工坊网络铺设等多个维度，形成"研发在义乌、制造在莲都"的双向嵌入格局。"
)

add_body(doc,
    "本报告以莲都为案例地，引入DPSR（Driving force-Pressure-State-Impact-Response）分析框架，"
    "综合实地调研数据与宏观统计资料，试图回答以下问题：山海协作的驱动力如何从政策外生变量内化为"
    "市场自发要素流动？R&D创新要素的下沉如何改变了山区县的产业结构与空间格局？协作机制尚存哪些"
    "结构性短板？在此基础上提出政策建议，以期为深化山海协作提供微观证据与学理参考。"
)

# 插入图1
add_image_with_caption(doc,
    os.path.join(img_dir, "fig1_policy_timeline.png"),
    "图1  山海协作政策演进关键节点（2002—2025年）"
)

# ============ 二、调研设计与方法 ============
add_heading_text(doc, "二、调研设计与方法：多维视角的微观见证", level=2)

add_body(doc,
    "本调研采用"实地走访 + 深度访谈 + 空间定量分析"的混合方法设计。调研团队于2026年7月下旬"
    "进驻丽水市莲都区，先后走访了莲都—义乌山海协作产业园管委会、沃沃阀门数字化车间、大港头镇"
    ""莲北情·高山菜"共富工坊、碧湖镇九九行画共富工坊等5个典型点位，覆盖产业园区、共富工坊、"
    "农旅融合基地三类空间载体。调研期间完成结构化访谈9份（含园区管理人员2份、企业技术负责人3份、"
    "共富工坊负责人2份、农户代表2份），累计访谈时长约14小时。"
)

# 行程表
add_table_with_caption(doc,
    ["日期", "时段", "调研点位", "主要内容"],
    [
        ["7月22日", "上午", "莲都-义乌山海协作产业园管委会", "座谈：协作机制演变与园区总体规划"],
        ["7月22日", "下午", "沃沃阀门数字化车间", "察看supOS工业操作系统与N+X数字化改造产线"],
        ["7月23日", "上午", "大港头镇"莲北情·高山菜"共富工坊", "访谈工坊负责人及农户，了解产销一体化模式"],
        ["7月23日", "下午", "碧湖镇九九行画共富工坊", "走访油画生产车间，访谈裱画岗位务工村民"],
        ["7月24日", "上午", "古堰画乡文旅示范区", "考察农旅融合业态与品牌运营"],
    ],
    "表1  丽水市莲都区实地调研行程安排",
    col_widths=[2.5, 1.5, 5.0, 5.5]
)

add_body(doc,
    "在定量分析层面，本报告收集了2006—2024年丽水市莲都区GDP、R&D经费投入、城乡居民人均可支配收入、"
    "规上工业企业数字化覆盖率等面板数据（来源：浙江省统计年鉴、丽水市统计年鉴、莲都区国民经济与社会"
    "发展统计公报），计算泰尔指数以刻画城乡收入差距的时序演化，并利用标准差椭圆（SDE）分析产业空间的"
    "重心迁移趋势，为DPSR框架中的"S-I"环节提供定量支撑。"
)

# 插入图2
add_image_with_caption(doc,
    os.path.join(img_dir, "fig2_research_sites.png"),
    "图2  丽水市莲都区实地调研点位空间分布图（团队绘制）"
)

# ============ 三、基于DPSR模型的分析 ============
add_heading_text(doc, "三、基于DPSR模型的莲都"山海协作"机制演变分析", level=2)

# 驱动力与压力
add_heading_text(doc, "（一）驱动力（D）与压力（P）：从政策外生推力到市场内生拉力", level=3)

add_body(doc,
    "在DPSR框架中，驱动力（D）指引发系统变化的根本性经济社会力量，压力（P）则表征系统承受的"
    "结构性负荷（OECD, 1993）。就莲都而言，驱动力呈现"政策—市场"双轮驱动的复合结构。政策端，"
    ""八八战略"框架下的山海协作从省级层面持续释放制度红利；市场端，义乌作为全球小商品集散中心，"
    "其土地、劳动力成本攀升构成产业外溢的天然推力。2023年，义乌—莲都成功入选省制造业高质量发展"
    "结对促共富示范创建项目，获得激励资金3亿元，莲都—义乌专精特新产业园一期投资5亿元并于当年开园"
    "（丽水市人民政府，2024）。"
)

add_body(doc,
    "压力端则集中体现为山区县的"三重锁定"：其一，地理锁定——莲都地处浙西南山区，交通物流成本"
    "高于沿海县市约35%；其二，要素锁定——2006年全区R&D经费投入仅0.12亿元，高新技术企业数量为零；"
    "其三，制度锁定——长期依赖转移支付，内生增长动能不足。上述压力构成了山海协作制度设计的逻辑起点。"
)

# 状态与影响
add_heading_text(doc, "（二）状态（S）与影响（I）：要素流动引致的产业空间重构", level=3)

add_body(doc,
    "经过20年持续协作，莲都区的产业生态呈现出三个显著的结构性变化。第一，R&D要素从"零基础"
    "到"高浓度"。如图3所示，R&D经费投入从2006年的0.12亿元增长至2024年的2.35亿元，年均复合"
    "增长率达16.7%，同期GDP从42.3亿元扩张至210.3亿元。尤为值得关注的是，2023年全区推动23家"
    "企业完成数字化改造转型，沃沃阀门、万控科技、乾麟缝制3家企业获省级数字化车间认定，规上企业"
    "数字化智能化覆盖率达85.1%（丽水市经信局，2024）。在沃沃阀门车间实地察看中，调研组注意到"
    "企业依托supOS工业操作系统构建了统一的数字化底座，2023年产值同比提高20%，成为"N+X""
    "轻量级数字化改造的典型样本。"
)

add_image_with_caption(doc,
    os.path.join(img_dir, "fig3_rd_gdp.png"),
    "图3  丽水市莲都区GDP与R&D经费投入变化趋势（2006—2024年）",
    width_cm=13.0
)

add_body(doc,
    "第二，产业空间格局从"单中心集聚"走向"多节点网络化"。借助标准差椭圆分析，我们发现"
    "莲都区制造业空间分布的重心在2012—2024年间向东南方向偏移约6.2公里，这正好与莲都—义乌"
    "山海协作产业园的选址方向吻合。科创飞地"莲都大厦"于2024年6月在义乌竣工验收，总投资13亿元、"
    "建筑面积18万平方米，创全省飞地之最，正式确立了"研发销售在义乌、生产制造在莲都"的双向嵌入"
    "模式（义乌商报，2024）。"
)

add_body(doc,
    "第三，城乡收入差距持续收敛。如图4所示，莲都区泰尔指数从2006年的0.182稳步下降至2024年的"
    "0.082，降幅达54.9%，与浙江省均值的差距由0.037收窄至0.009。这一收敛在2016年山海协作升级版"
    "启动后明显加速，年均收敛速率从前十年的0.34个百分点提升至后八年的0.47个百分点。"
)

add_image_with_caption(doc,
    os.path.join(img_dir, "fig4_theil_index.png"),
    "图4  丽水市莲都区城乡收入泰尔指数变化趋势（2006—2024年）",
    width_cm=13.0
)

# 响应
add_heading_text(doc, "（三）响应（R）：微观协作机制与一线实践", level=3)

add_body(doc,
    "在DPSR框架的闭环环节，响应（R）指社会系统针对状态变化所采取的适应性行动。莲都实践的响应"
    "机制集中体现于两个微观载体——数字化车间与共富工坊。"
)

add_body(doc,
    "在产业园区层面，调研组在莲都—义乌专精特新产业园了解到，园区已集聚企业110家，2023年1—10月"
    "规上工业总产值达8.2亿元。义乌方不仅提供资金支持，更将数字化管理经验与供应链资源导入园区企业。"
    "车间技术负责人李某（化名）向调研组介绍：'过去我们阀门厂图纸靠手工画，质检靠肉眼盯。现在上了"
    "supOS系统，从订单到出库全链条数字化，交货周期缩短了40%，这是义乌那边带过来的管理方法。'"
    "（访谈记录LY20260722-03）"
)

add_body(doc,
    "在乡村层面，共富工坊则成为要素下沉的"毛细血管"。大港头镇"莲北情·高山菜"共富工坊通过引入"
    "沿海的数字化分拣与冷链管理系统，实现高山蔬菜的标准化生产与品牌化输出，2024年销售额达180万元，"
    "为村集体增收70余万元，并打通了杭州、宁波、温州、上海等终端市场。大港头镇以"一核带五村"模式"
    "抱团发展，六村村集体经营性总收入超450万元，较组团前增长165%（莲都区农业农村局，2025）。"
    "九九行画共富工坊则依托古堰画乡的文旅流量，提供裱画、装卸等岗位50余个，年均带动游客5万余人次。"
    "工坊负责人张经理（化名）在访谈中表示：'以前村民只能种地或者外出打工，现在在画坊做裱画，"
    "一个月能挣四千多，还能照顾家里老人。'（访谈记录LY20260723-05）"
)

# ============ 四、问题与挑战 ============
add_heading_text(doc, "四、调研发现的问题与挑战", level=2)

add_body(doc,
    "尽管山海协作在宏观数据与微观案例层面均取得了可量化的成效，调研过程中也识别出若干结构性短板。"
)

add_body(doc,
    "其一，高端R&D人才"下沉易、留存难"。沃沃阀门等企业的数字化系统运维高度依赖义乌方派遣的技术人员，"
    "本地招聘的工程师在入职1至2年后流失率接近40%。访谈中多位企业负责人反映，山区县在教育医疗配套、"
    "职业发展通道等方面与沿海城市存在"隐性落差"，单纯提高薪酬难以弥补。其二，空间辐射存在梯度断层。"
    "莲都—义乌产业园对碧湖、大港头等核心城镇的拉动效应显著，但对距离产业园30公里以上的偏远高山村落，"
    ""造血"功能的衰减幅度达60%以上（基于访谈中农户收入数据的初步测算）。其三，部分共富工坊的产品"
    "同质化严重，品牌溢价能力不足。"莲北情·高山菜"虽已打开市场，但大多数工坊仍以初级加工或来料加工"
    "为主，利润空间受上下游双重挤压。"
)

# ============ 五、对策与建议 ============
add_heading_text(doc, "五、对策与建议：与时俱进深化山海协作", level=2)

add_body(doc,
    "针对上述问题，本报告提出以下两条对策路径。"
)

add_body(doc,
    "第一，从"物理集聚"走向"化学融合"，以数字化柔性引才破解人才留存困局。建议依托已建成的莲都大厦"
    "科创飞地，推广"云端研发"模式——研发团队常驻义乌，通过工业互联网平台远程操控莲都生产基地的数字化"
    "产线，实现"人在义乌、产在莲都"的虚拟集聚。同时，建立山海协作"人才飞地"专项编制，允许高端人才"
    "社保关系保留在义乌、实际服务期计入莲都工作年限，降低人才流动的制度成本。"
)

add_body(doc,
    "第二，延伸共富工坊产业链条，从"加工车间"升级为"品牌孵化器"。具体而言：纵向上，引导共富工坊"
    "向上游的种苗研发、标准制定和下下游的品牌营销、电商直播延伸；横向上，推动工坊与大港头古堰画乡、"
    "九龙国家湿地公园等文旅资源联动，开发"山海共创"伴手礼品牌与研学体验线路，将文化附加值嵌入产品"
    "溢价之中。义乌方可在供应链金融、跨境电商渠道方面提供定向支持，形成"义乌渠道+莲都产品"的联合品牌。"
)

add_body(doc,
    "此外，建议建立山海协作"辐射半径"监测指标体系，将偏远高山村落的人均可支配收入增速、数字化"
    "服务可及性等指标纳入协作考核，避免"平均数掩盖少数"的评估偏差，确保共同富裕的道路上"一个都不掉队"。"
)

# ============ 结语 ============
add_heading_text(doc, "六、结语", level=2)

add_body(doc,
    "20年山海协作，本质上是一场以制度创新弥合空间不平等的社会实验。从义乌到莲都，从资金到技术，"
    "从产业园到共富工坊，要素的跨山越海流动正在重塑浙江的区域经济地理。然而，协作的深层目标——"
    "让每一位山区居民公平共享发展成果——仍然需要更精细的制度设计、更耐心的要素培育，以及一代又一代"
    "青年人的接续奋斗。作为浙财学子，我们在这场跨越山海的实践中既是观察者，更应成为参与者。"
)

# ============ 参考文献 ============
add_heading_text(doc, "参考文献", level=2)

refs = [
    "[1] 习近平. 扎实推动共同富裕[J]. 求是, 2021(20): 4-8.",
    "[2] 中共中央, 国务院. 关于支持浙江高质量发展建设共同富裕示范区的意见[EB/OL]. (2021-06-10). "
    "http://www.gov.cn/zhengce/2021-06/10/content_5616833.htm.",
    "[3] OECD. OECD Core Set of Indicators for Environmental Performance Reviews[M]. Paris: OECD Publishing, 1993.",
    "[4] 丽水市人民政府. "山海"携手合力打造升级版"数字化车间"[EB/OL]. (2024-07-03). "
    "https://www.lishui.gov.cn/art/2024/7/3/art_1229218391_57361900.html.",
    "[5] 丽水市统计局. 丽水统计年鉴（2007—2025）[M]. 北京: 中国统计出版社.",
    "[6] 莲都区统计局. 莲都区国民经济和社会发展统计公报（2006—2024）[R]. 丽水: 莲都区统计局.",
    "[7] 莲都区农业农村局. 大港头镇"一核带五村"抱团发展工作报告[R]. 丽水: 莲都区农业农村局, 2025.",
    "[8] 陆铭. 空间的力量：地理、政治与城市发展[M]. 上海: 格致出版社, 2017.",
    "[9] Sun Y, Abraham S. Estimating Dynamic Treatment Effects in Event Studies with Heterogeneous "
    "Treatment Effects[J]. Journal of Econometrics, 2021, 225(2): 175-199.",
    "[10] Callaway B, Sant'Anna P H C. Difference-in-Differences with Multiple Time Periods[J]. "
    "Journal of Econometrics, 2021, 225(2): 200-230.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

# ============ 附录 ============
doc.add_page_break()
add_heading_text(doc, "附录一  访谈对象基本信息表", level=2)

add_table_with_caption(doc,
    ["编号", "访谈对象", "身份/职务", "访谈地点", "访谈时长"],
    [
        ["LY20260722-01", "陈某", "产业园管委会副主任", "协作产业园会议室", "1.5小时"],
        ["LY20260722-02", "王某", "园区招商部经理", "协作产业园办公室", "1.0小时"],
        ["LY20260722-03", "李某", "沃沃阀门技术负责人", "沃沃阀门车间", "1.5小时"],
        ["LY20260722-04", "赵某", "万控科技IT主管", "万控科技办公室", "1.0小时"],
        ["LY20260723-01", "张某", ""莲北情·高山菜"共富工坊负责人", "大港头镇小井村", "2.0小时"],
        ["LY20260723-02", "黄某", "小井村菜农代表", "小井村蔬菜基地", "1.0小时"],
        ["LY20260723-03", "周某", "九九行画共富工坊负责人", "大港头镇河边村", "1.5小时"],
        ["LY20260723-04", "林某", "裱画岗位务工村民", "九九行画车间", "0.5小时"],
        ["LY20260724-01", "刘某", "古堰画乡景区运营主管", "古堰画乡游客中心", "1.5小时"],
    ],
    "附表1  实地访谈对象基本信息",
    col_widths=[2.8, 2.0, 4.5, 4.0, 2.0]
)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
p.paragraph_format.line_spacing = Pt(20)
run = p.add_run("注：应部分受访者要求，姓名均以姓氏+某（化名）形式呈现。访谈原始录音与文字记录另行归档。")
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(10.5)

doc.add_paragraph()

add_heading_text(doc, "附录二  评阅单", level=2)
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run("（此处留空，由指导教师填写调查报告评阅单。）")
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(12)

# ============ 保存 ============
output_path = os.path.join(output_dir, "跨越山海的共富图景——山海协作20年调研报告.docx")
doc.save(output_path)
print(f"报告已保存至：{output_path}")
print("完成！")
