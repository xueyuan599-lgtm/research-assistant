"""生成 FEM 泊松方程求解报告 (中文版 Word)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

# ── 封面标题 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.space_before = Pt(120)
run = title.add_run("有限元方法\n泊松方程数值求解报告")
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("基于 scikit-fem 的 P1 有限元实现")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 100, 100)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(40)
run = p.add_run("2026-07-02")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ── 1. 问题描述 ──
doc.add_heading("1. 问题描述", level=1)
doc.add_paragraph(
    "在单位正方形区域上考虑带有齐次 Dirichlet 边界条件的泊松方程：",
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(r"-∇²u = 8π² sin(2πx) sin(2πy),  (x,y) ∈ (0,1)²")
run.font.size = Pt(12)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(r"u = 0  在边界 ∂Ω 上")
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph(
    "该问题具有已知的解析解："
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(r"u_exact(x,y) = sin(2πx) sin(2πy)")
run.font.size = Pt(12)
run.italic = True

# ── 2. 数值方法 ──
doc.add_heading("2. 数值方法", level=1)
doc.add_paragraph("空间离散：", style="List Bullet")
doc.add_paragraph("采用连续 Galerkin (CG) 有限元方法，使用线性三角形单元（P1）。")
doc.add_paragraph("网格生成：", style="List Bullet")
doc.add_paragraph("通过 15×15 张量积网格点生成结构化三角形网格，并经过 2 次均匀加密。")
doc.add_paragraph("线性求解器：", style="List Bullet")
doc.add_paragraph("直接稀疏求解器（scipy.sparse.linalg.spsolve）。")
doc.add_paragraph("实现工具：", style="List Bullet")
doc.add_paragraph("scikit-fem（轻量级纯 Python 有限元库）。")

# ── 3. 结果 ──
doc.add_heading("3. 计算结果", level=1)

fig_path = "E:/wuyi/数学建模半自动/research-assistant/outputs/fem_poisson_demo.png"
if os.path.exists(fig_path):
    doc.add_picture(fig_path, width=Inches(6.0))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("图 1: 数值解（P1 有限元）、精确解与逐点绝对误差对比")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_heading("3.1 误差分析", level=2)
table = doc.add_table(rows=3, cols=2, style="Light Shading Accent 1")
table.cell(0, 0).text = "指标"
table.cell(0, 1).text = "数值"
table.cell(1, 0).text = "相对 L2 误差"
table.cell(1, 1).text = "0.14%"
table.cell(2, 0).text = "视觉评估"
table.cell(2, 1).text = "数值解与精确解在视觉上几乎无法区分"

# ── 4. 讨论 ──
doc.add_heading("4. 分析与讨论", level=1)
doc.add_paragraph(
    "P1 有限元在中等加密程度网格上取得了很高的求解精度（相对 L2 误差 = 0.14%）。"
    "误差主要集中在梯度较大的区域附近，"
    "可通过自适应网格加密（h-加密）或使用高阶单元（p-加密）进一步降低。"
)
doc.add_paragraph(
    "scikit-fem 库提供了轻量级的纯 Python 有限元框架，"
    "适合 PDE 求解器的快速原型开发、算法验证及自动化科研管线集成。"
    "相比完整的 FEniCS 框架，scikit-fem 安装简便（pip install），"
    "在 Windows 环境下具有良好的兼容性。"
)

# ── 5. 复现说明 ──
doc.add_heading("5. 复现说明", level=1)
doc.add_paragraph("在终端中执行以下命令即可复现结果：")
p = doc.add_paragraph()
run = p.add_run("python scripts/fem_demo.py")
run.font.name = "Consolas"
run.font.size = Pt(10)

doc.add_paragraph("")
table2 = doc.add_table(rows=4, cols=2, style="Light Shading Accent 1")
table2.cell(0, 0).text = "依赖项"
table2.cell(0, 1).text = "说明"
table2.cell(1, 0).text = "scikit-fem"
table2.cell(1, 1).text = "有限元核心库"
table2.cell(2, 0).text = "numpy / scipy"
table2.cell(2, 1).text = "数值计算与稀疏求解"
table2.cell(3, 0).text = "matplotlib"
table2.cell(3, 1).text = "可视化与图表生成"

out_path = "E:/wuyi/数学建模半自动/research-assistant/outputs/fem-demo/fem_report.docx"
doc.save(out_path)
print("OK")
